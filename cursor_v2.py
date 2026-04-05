"""
ATS Resume Builder v2 — Groq dual-model pipeline (extraction + assembler), 8 LangChain steps + Streamlit
=====================================================================
Wave 1 (parallel): Keywords, Skills Gap
Wave 2 (parallel): Summary, Skills, Projects, Experience
Wave 3 (parallel): Cold Email, Cover Letter

Improvements over v1:
- Education section with bold institution names and right-aligned dates
- Enhanced prompts targeting >90% ATS score on Jobscan
- Triple keyword coverage (summary + skills + bullets)
- Post-assembly keyword verification & injection
- Visual formatting matched to reference layout

Run:
  pip install streamlit langchain-groq langchain-core python-docx python-dotenv
  set GROQ_API_KEY=your_key_here
  Optional: set GROQ_EXTRACTION_MODEL and GROQ_ASSEMBLER_MODEL (or GROQ_MODEL as default for both).
  streamlit run cursor_v2.py
"""

from __future__ import annotations

import io
import json
import os
import re
import time
import zipfile
from datetime import datetime
from typing import Dict, List, Tuple

from dotenv import load_dotenv
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

load_dotenv()

import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableParallel  # kept for reference

# ── Config ────────────────────────────────────────────────────────────────────

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "").strip()
_DEFAULT_GROQ_MODEL = "moonshotai/kimi-k2-instruct"
# GROQ_MODEL: default when GROQ_EXTRACTION_MODEL / GROQ_ASSEMBLER_MODEL are not set
GROQ_MODEL = os.environ.get("GROQ_MODEL", _DEFAULT_GROQ_MODEL).strip()
# Extraction: ATS keywords + skills gap (steps 1–2). Assembler: summary → cover (steps 3–8).
GROQ_EXTRACTION_MODEL = os.environ.get("GROQ_EXTRACTION_MODEL", GROQ_MODEL).strip()
GROQ_ASSEMBLER_MODEL = os.environ.get("GROQ_ASSEMBLER_MODEL", GROQ_MODEL).strip()

# Sidebar catalog (IDs must match Groq Cloud). Unknown env models are prepended in the UI.
GROQ_MODEL_CATALOG: List[str] = [
    "moonshotai/kimi-k2-instruct",
    "meta-llama/llama-3.3-70b-versatile",
    "meta-llama/llama-3.1-8b-instant",
    "openai/gpt-oss-120b",
]


def _groq_select_options(current: str) -> List[str]:
    opts = list(GROQ_MODEL_CATALOG)
    if current and current not in opts:
        opts.insert(0, current)
    return opts


def _groq_select_index(options: List[str], current: str) -> int:
    try:
        return options.index(current)
    except ValueError:
        return 0


C_BLACK = RGBColor(0x00, 0x00, 0x00)
C_DARK = RGBColor(0x1A, 0x1A, 0x1A)
C_GRAY = RGBColor(0x33, 0x33, 0x33)

# ── Hardcoded links — applied to every generated resume ──────────────────────

LINK_MAP: Dict[str, str] = {
    "linkedin": "https://www.linkedin.com/in/sai-harshith-kolavasi-69459a243/",
    "github": "https://github.com/SaiHarshithK22",
    "email": "harshith.kolavasi@gmail.com",
    "research rag application": "https://saiharshithk22-research-rag-application-main-50zvlk.streamlit.app/",
    "vehicle damage detection system": "https://saiharshithk22-vehicle-damage-detection-app-7w2rdy.streamlit.app/",
    "health insurance premium prediction system": "https://saiharshithk22-ml-health-insurance-premium-predicti-main-3zcf0v.streamlit.app/",
    "credit risk modelling": "https://saiharshithk22-credit-risk-main-uzbvcn.streamlit.app/",
}

# ── Section parsing constants ─────────────────────────────────────────────────

SECTION_HEADERS = {
    "PROFESSIONAL SUMMARY", "SUMMARY PROFESSIONAL", "SUMMARY", "OBJECTIVE", "PROFILE",
    "TECHNICAL SKILLS", "SKILLS", "CORE SKILLS", "KEY SKILLS",
    "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EXPERIENCE",
    "EMPLOYMENT", "RELEVANT EXPERIENCE", "CAREER HISTORY",
    "PROJECTS", "SELECTED PROJECTS", "KEY PROJECTS", "RELEVANT PROJECTS",
    "PERSONAL PROJECTS", "ACADEMIC PROJECTS",
    "EDUCATION", "ACADEMIC BACKGROUND",
    "CERTIFICATIONS", "LICENSES AND CERTIFICATIONS",
    "ACHIEVEMENTS", "AWARDS", "HONORS", "PUBLICATIONS",
    "LANGUAGES", "VOLUNTEER", "VOLUNTEER EXPERIENCE",
}

PROJECT_HEADERS = {
    "PROJECTS", "SELECTED PROJECTS", "KEY PROJECTS",
    "RELEVANT PROJECTS", "PERSONAL PROJECTS", "ACADEMIC PROJECTS",
}

EXPERIENCE_HEADERS = {
    "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EXPERIENCE",
    "EMPLOYMENT", "CAREER HISTORY",
}

RELEVANT_EXP_HEADERS = {
    "RELEVANT EXPERIENCE",
}

DATE_RE = re.compile(
    r"\b((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}"
    r"\s*[–—-]\s*(?:Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}))\b",
    re.IGNORECASE,
)

INSTITUTION_KEYWORDS = {
    "university", "college", "institute", "school", "academy",
    "polytechnic", "conservatory",
}

DEGREE_KEYWORDS = {
    "bachelor", "master", "phd", "doctor", "associate",
    "b.s.", "m.s.", "b.a.", "m.a.", "mba", "diploma",
    "master of science", "master of arts", "bachelor of science",
    "bachelor of arts", "doctor of philosophy",
}


# ── Text helpers ──────────────────────────────────────────────────────────────

def safe_json(text: str) -> dict:
    text = re.sub(r"```json|```", "", text).strip()
    try:
        return json.loads(text)
    except Exception:
        return {}


def extract_text_from_docx(uploaded_file) -> str:
    doc = DocxDocument(io.BytesIO(uploaded_file.read()))
    lines: List[str] = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in lines:
                    lines.append(t)
    return "\n".join(lines)


def extract_name_contact(resume: str) -> Tuple[str, str]:
    lines = [l.strip() for l in resume.splitlines() if l.strip()]
    if not lines:
        return "Your Name", ""
    name = lines[0]
    contact = ""
    for line in lines[1:8]:
        if "|" in line or "@" in line or re.search(r"\d{3}[-.\s]?\d{3}[-.\s]?\d{4}", line):
            contact = line
            break
    return name, contact


def _extract_section(resume: str, header_set: set) -> str:
    lines = resume.splitlines()
    start = None
    for i, line in enumerate(lines):
        if line.upper().strip().rstrip(":") in header_set:
            start = i
            break
    if start is None:
        return ""
    end = len(lines)
    for j in range(start + 1, len(lines)):
        u = lines[j].upper().strip().rstrip(":")
        if u in SECTION_HEADERS and u not in header_set:
            end = j
            break
    return "\n".join(lines[start:end]).strip()


def extract_projects(resume: str) -> str:
    return _extract_section(resume, PROJECT_HEADERS)


def extract_experience(resume: str) -> str:
    return _extract_section(resume, EXPERIENCE_HEADERS)


def extract_relevant_experience(resume: str) -> str:
    return _extract_section(resume, RELEVANT_EXP_HEADERS)


def extract_education(resume: str) -> str:
    return _extract_section(resume, {"EDUCATION", "ACADEMIC BACKGROUND"})


def normalize(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\t", " ")
    out: List[str] = []
    prev_blank = False
    for raw in text.splitlines():
        line = re.sub(r"[ ]{2,}", " ", raw).strip()
        line = re.sub(r"^[-*•·▪◦]\s*", "• ", line)
        if not line:
            if not prev_blank:
                out.append("")
            prev_blank = True
            continue
        out.append(line)
        prev_blank = False
    text = "\n".join(out).strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def dedupe_projects(text: str) -> str:
    if not text:
        return ""
    lines = text.splitlines()
    out: List[str] = []
    seen: set = set()
    skip = False
    for line in lines:
        s = line.strip()
        if not s:
            if not skip:
                out.append("")
            continue
        if s.upper().rstrip(":") in PROJECT_HEADERS:
            if any(l.strip().upper().rstrip(":") in PROJECT_HEADERS for l in out):
                continue
            out.append(s)
            continue
        is_proj_header = ("|" in s) and bool(DATE_RE.search(s))
        if is_proj_header:
            key = re.sub(r"\s+", " ", s.lower())
            if key in seen:
                skip = True
                continue
            seen.add(key)
            skip = False
        if skip:
            continue
        out.append(s)
    return normalize("\n".join(out))


def fix_project_bullets(text: str) -> str:
    """Ensure every non-header line in project text starts with a bullet character."""
    if not text:
        return ""
    lines = text.splitlines()
    out: List[str] = []
    for line in lines:
        s = line.strip()
        if not s:
            out.append("")
            continue
        is_header = ("|" in s and bool(DATE_RE.search(s))) or s.upper().rstrip(":") in SECTION_HEADERS
        is_already_bullet = bool(re.match(r"^[•·\-*▪◦]", s))
        if not is_header and not is_already_bullet and not s.startswith("|"):
            s = "• " + s
        out.append(s)
    return "\n".join(out)


def _is_institution_line(line: str) -> bool:
    low = line.lower()
    return any(k in low for k in INSTITUTION_KEYWORDS)


def _is_degree_line(line: str) -> bool:
    low = line.lower()
    return any(k in low for k in DEGREE_KEYWORDS)


# ═══════════════════════════════════════════════════════════════════════════════
# PROMPTS  (v2 — balanced keyword matching for >90% Jobscan score)
# ═══════════════════════════════════════════════════════════════════════════════

# ── Wave 1: Keyword + Gap analysis ──────────────────────────────────────────

PROMPT_KEYWORDS = """You are an expert ATS keyword analyst. ATS systems do EXACT string matching.

Extract keywords from the Job Description and compare against the resume.

═══ WHAT COUNTS AS A KEYWORD ═══
YES — extract these:
- Programming languages (Python, Java, Go, C++)
- Named tools, frameworks, libraries (LangChain, PyTorch, FastAPI, Docker, Terraform)
- Named platforms and cloud services (AWS, S3, SageMaker, ECS, Lambda, Databricks)
- Specific methodologies (RAG, CI/CD, Agile, prompt engineering, fine-tuning)
- ML/AI concepts (anomaly detection, predictive modelling, time-series forecasting, LLMs)
- Named AI models (OpenAI, Gemini, Llama, Claude, Qwen)
- Technical terms (function calling, vector retrieval, model deployment, containerized services)
- Soft skills explicitly stated (problem-solving, communication, collaboration, ownership)
- Job title as stated in the JD

NO — do NOT extract these:
- Hours/schedule ("20 hours per week", "full-time")
- Degree names ("Computer Science", "Applied Math", "PhD")
- Job roles ("data analyst", "business owner", "hiring manager")
- Generic business phrases ("high-impact projects", "deployable artifacts", "handoff docs")
- Vague activity descriptions ("present technical work", "building AI prototypes")
- Company-specific context ("Goldman Sachs", "meritocracy")

Return ONLY valid JSON (no markdown fences):
{{
  "job_title": "exact job title from JD",
  "jd_keywords": ["all real keywords — tools, technologies, methodologies, skills"],
  "matched": ["keywords already in resume"],
  "missing": ["keywords NOT in resume that should be added"],
  "for_summary": ["top 8-10 technical keywords for the summary"],
  "for_skills": ["tools, languages, frameworks, platforms, methodologies, and soft skills ONLY"],
  "for_bullets": ["technical terms and action verbs to use in project bullets"]
}}

JOB DESCRIPTION:
{jd}

RESUME:
{resume}
"""

PROMPT_SKILLS_GAP = """You are a senior technical recruiter performing a skills gap analysis.

Return ONLY valid JSON (no markdown fences):
{{
  "required_skills": ["skill"],
  "candidate_has": ["skill"],
  "gaps": [
    {{"skill": "", "importance": "Critical|Important|Nice-to-have", "suggestion": "how to address in resume"}}
  ],
  "transferable": ["candidate skill that substitutes for a gap"]
}}

JOB DESCRIPTION:
{jd}

RESUME:
{resume}
"""

# ── Wave 2: Content agents (receive keyword analysis) ────────────────────────

PROMPT_SUMMARY = """You are an expert resume writer maximizing ATS keyword match rate.

Rewrite the professional summary so an ATS scanner finds JD keywords in it.

KEYWORDS TO EMBED (use exact phrasing):
{for_summary}

Rules:
1. First sentence MUST start with: "{job_title} with X+ years of hands-on experience in"
   followed by the top 3-4 JD keywords.
2. Use the EXACT JD phrasing, not synonyms. If JD says "large-scale", write "large-scale".
3. Reference 2-3 quantified achievements from the resume (never invent metrics).
4. Write 3-4 sentences, one paragraph, no bullets.
5. No first person ("I am"), no clichés ("passionate", "hard-working").
6. Each sentence should contain 2-3 JD keywords, woven in naturally — the summary must read
   like a human wrote it, not like a keyword list.

Return ONLY valid JSON (no markdown fences):
{{
  "optimized_summary": "rewritten summary (3-4 sentences)",
  "keywords_embedded": ["JD keywords you used"]
}}

JOB DESCRIPTION:
{jd}

RESUME:
{resume}
"""

PROMPT_SKILLS = """You are an ATS resume writer building a Technical Skills section.

JD KEYWORDS TO INCLUDE:
{for_skills}

═══ WHAT BELONGS IN THIS SECTION ═══
INCLUDE ONLY:
- Programming languages (Python, SQL, Java, Go, C++)
- Named ML/DL libraries and algorithms (scikit-learn, PyTorch, XGBoost, ResNet, SMOTE)
- Named AI/LLM tools and concepts (LangChain, RAG, FAISS, ChromaDB, prompt engineering,
  function calling, vector retrieval, embeddings, fine-tuning, agent architectures)
- Named AI models (OpenAI, Gemini, Llama, Qwen, Claude)
- Named frameworks and dev tools (FastAPI, Streamlit, Git, Docker, REST APIs)
- Named cloud services (AWS, S3, EC2, Lambda, SageMaker, ECS/EKS, DynamoDB, Terraform)
- Named data platforms (Databricks, Redshift)
- Named visualization tools (Power BI, Tableau, Matplotlib, Seaborn)
- Real soft skills from JD (problem-solving, communication, ownership, collaboration)

NEVER INCLUDE — these are NOT skills:
- Generic JD phrases ("deployable artifacts", "handoff docs", "impact metrics", "runbooks")
- Job roles ("data analysts", "business owners", "nontechnical stakeholders")
- Degree fields ("Computer Science", "Applied Math", "PhD candidate")
- Hours or logistics ("20 hours per week")
- Vague activities ("building prototypes", "present technical work", "harden prototypes")
- Company names or products

Rules:
- Format: "Category: item1, item2, item3" (one row per category).
- Use the JD's EXACT spellings.
- Group into: Languages, ML / DL, Generative AI / LLMs, Frameworks / Tools,
  Cloud / Infrastructure, Analytics / Visualization, Soft Skills (if JD mentions any).
- Each category row should have real, recognizable skills — not sentences or phrases.
- Keep each row concise and scannable.

Return ONLY valid JSON (no markdown fences):
{{
  "skills_lines": [
    "Languages: ...",
    "ML / DL: ...",
    "Generative AI / LLMs: ...",
    "Frameworks / Tools: ...",
    "Cloud / Infrastructure: ...",
    "Analytics / Visualization: ...",
    "Soft Skills: ... (only if JD has soft skills)"
  ]
}}

JOB DESCRIPTION:
{jd}

RESUME:
{resume}
"""

PROMPT_PROJECTS = """You are an ATS resume writer tailoring the PROJECTS section to THIS job description.

TARGET ROLE (from JD analysis):
{job_title}

JD KEYWORDS AVAILABLE:
{for_bullets}

MISSING KEYWORDS (not yet in resume — add where they naturally fit):
{missing_keywords}

═══ JD ALIGNMENT (PRIMARY) ═══
- Order projects by relevance to THIS JD: put the project that best matches the JD's tools,
  domain, and responsibilities FIRST; less relevant projects follow.
- For each project, lead with bullets that map to the JD's top requirements (read the full JD).
- Rewrite bullets so outcomes and tech explicitly echo JD language where truthful — same work,
  stronger overlap with what the employer asked for.
- If the JD stresses a stack (e.g. cloud, ML, LLMs, data pipelines), foreground that stack in
  the matching project's bullets first.

═══ CRITICAL FORMAT ═══
For EACH project, output EXACTLY this (header all on one line):
ProjectName | TechStack | Link                DateRange
• bullet 1
• bullet 2
• bullet 3
• bullet 4
• bullet 5

═══ RULES ═══
1. Keep project names, tech stacks, and date ranges EXACTLY as in the resume.
2. Start EVERY bullet with a strong action verb.
3. Each bullet should include 1-2 JD keywords WHERE THEY NATURALLY FIT.
   Do NOT force keywords that make the bullet unreadable.
4. PRESERVE the original meaning of each bullet — describe what the project ACTUALLY does.
   Do NOT fabricate capabilities the project doesn't have.
5. Preserve ALL original metrics, numbers, percentages. Never invent metrics.
6. Each project: 4-5 bullets. Rewrite existing bullets; do not add imaginary ones.
7. Do NOT include Education or Certifications.
8. Include "| Link" after tech stack for projects with deployable links.
9. The bullets must read naturally — a human reviewer should not notice keyword stuffing.
10. NEVER add phrases like "20 hours per week", "handoff docs", "business owners",
    "nontechnical stakeholders", "deployable artifacts" into bullets — these are not technical work.

Return ONLY valid JSON (no markdown fences):
{{
  "projects_text": "full plain-text PROJECTS block following the format above"
}}

JOB DESCRIPTION:
{jd}

RESUME:
{resume}
"""

PROMPT_EXPERIENCE = """You are an ATS resume writer improving work experience bullets with JD keywords.

JD KEYWORDS AVAILABLE:
{for_bullets}

MISSING KEYWORDS (not yet in resume — add where they naturally fit):
{missing_keywords}

═══ CRITICAL FORMAT ═══
For EACH role, output EXACTLY this (header all on one line):
Job Title | Company Name                DateRange
• bullet 1
• bullet 2
• bullet 3

═══ RULES ═══
1. Keep job titles, company names, locations, and date ranges EXACTLY as in the resume.
2. Start EVERY bullet with a strong action verb.
3. Each bullet should include 1-2 JD keywords WHERE THEY NATURALLY FIT.
4. PRESERVE the original meaning — describe what the candidate ACTUALLY did.
5. Preserve ALL original metrics, numbers, percentages. Never invent metrics.
6. Each role: 3-5 bullets. Rewrite existing bullets; do not add imaginary duties.
7. If the original resume has NO work experience at all, return an empty string for experience_text.
8. Bullets must read naturally — no keyword stuffing.

Return ONLY valid JSON (no markdown fences):
{{
  "experience_text": "full plain-text WORK EXPERIENCE block following the format above"
}}

JOB DESCRIPTION:
{jd}

RESUME:
{resume}
"""

PROMPT_RELEVANT_EXP = """You are an ATS resume writer tailoring the RELEVANT EXPERIENCE section to THIS job description.

The RELEVANT EXPERIENCE section is a concise summary of hands-on work that should read as direct
evidence the candidate can perform THIS role.

TARGET ROLE (from JD analysis):
{job_title}

JD KEYWORDS AVAILABLE:
{for_bullets}

MISSING KEYWORDS (not yet in resume — add where they naturally fit):
{missing_keywords}

═══ JD ALIGNMENT (PRIMARY) ═══
- Read the full JD: identify must-have tools, responsibilities, and domain; every bullet should
  tie to at least one of them when the source material allows.
- ORDER bullets by strength of match to the JD (most compelling evidence for this job first).
- Reframe wording so responsibilities and outcomes mirror JD phrasing where truthful — do not
  invent employers, dates, or work you cannot infer from the resume.

═══ CRITICAL FORMAT ═══
Output EXACTLY this format (header all on one line):
Title | Organization                DateRange
• bullet 1
• bullet 2
• bullet 3
• bullet 4

═══ RULES ═══
1. Keep the title, organization, and date range EXACTLY as in the resume.
2. Start each bullet with a strong action verb (Built, Developed, Designed, Engineered, Deployed).
3. Include 1-2 JD keywords per bullet WHERE THEY NATURALLY FIT.
4. Preserve ALL original metrics, numbers, percentages. Never invent metrics.
5. Keep bullets concise — each bullet is a 1-2 line summary of a project.
6. 3-5 bullets total. Rewrite existing bullets; do not invent new ones.
7. If the original resume has NO relevant experience section, return an empty string.
8. Bullets must read naturally — no keyword stuffing.

Return ONLY valid JSON (no markdown fences):
{{
  "relevant_exp_text": "full plain-text RELEVANT EXPERIENCE block following the format above"
}}

JOB DESCRIPTION:
{jd}

RESUME:
{resume}
"""

# ── Wave 3: Email + Cover (uses assembled resume) ───────────────────────────

PROMPT_EMAIL = """You are an expert at writing cold outreach emails to recruiters and hiring managers.

IMPORTANT: The resume below is already ATS-optimized and tailored to the JD.
Use it as the definitive source. Never invent facts.

Rules:
1. Subject line: mention the exact job title + be curiosity-driving (not generic).
2. Opening: reference something specific from the company or JD — NOT "I saw your job posting".
3. Body: 2 short paragraphs max.
   - Para 1: who the candidate is + ONE most impressive quantified achievement from the resume
     that maps directly to the JD's top requirement.
   - Para 2: why THIS company specifically (reference their product, mission, or tech stack from JD).
4. CTA: one clear, low-friction ask (15-minute call, not "please review my resume").
5. Tone: confident, direct, human — NOT corporate or sycophantic.
6. Total length: under 150 words (excluding subject line).
7. Sign-off: candidate's name, phone, email, LinkedIn, GitHub from the resume header.

Return ONLY valid JSON (no markdown fences):
{{
  "subject_line": "...",
  "email_body": "full email text including sign-off"
}}

JOB DESCRIPTION:
{jd}

OPTIMIZED RESUME:
{optimized_resume}
"""

PROMPT_COVER = """You are a professional cover letter writer who specializes in ATS-optimized, human-readable letters.

IMPORTANT: The resume below is already ATS-optimized and tailored to the JD.
Use it as the definitive source. All keywords and achievements MUST come from this resume.

Rules:
1. Header: candidate's full contact info from the resume (Name, City, Phone, Email, LinkedIn, GitHub).
2. Date: use today's date.
3. Salutation: "Dear Hiring Manager," (unless a specific name is in the JD).
4. Paragraph 1 — Hook (2-3 sentences):
   - State the exact job title from the JD.
   - Lead with the candidate's single strongest quantified achievement from the resume
     that maps directly to the JD's top requirement.
   - Do NOT start with "I am writing to apply..."
5. Paragraph 2 — Fit (3-4 sentences):
   - Highlight 2-3 specific technical skills or project results from the resume.
   - Use EXACT keywords from the JD naturally.
   - Reference a specific project name and metric from the resume.
6. Paragraph 3 — Motivation (2-3 sentences):
   - Why THIS company specifically — reference their mission, product, or tech stack from the JD.
   - Connect the candidate's career direction to what they're building.
7. Closing paragraph (2 sentences):
   - Express enthusiasm, mention the enclosed resume.
   - Clear CTA: looking forward to discussing how you can contribute.
8. Sign-off: "Sincerely," followed by candidate's full name.
9. Tone: confident, specific, professional — no clichés like "passionate team player".
10. Total length: 300-380 words.

Return ONLY valid JSON (no markdown fences):
{{
  "cover_letter_text": "full formatted cover letter including header and sign-off"
}}

JOB DESCRIPTION:
{jd}

OPTIMIZED RESUME:
{optimized_resume}
"""


# ── Assembler ─────────────────────────────────────────────────────────────────

def _inject_missing_keywords(skills_block: str, missing: List[str], resume: str) -> str:
    """Inject any keywords still missing from the skills block."""
    if not missing or not skills_block:
        return skills_block
    low_block = skills_block.lower()
    low_resume = resume.lower()
    still_missing = [
        kw for kw in missing
        if kw.lower() not in low_block and kw.lower() not in low_resume
    ]
    if not still_missing:
        return skills_block

    lines = skills_block.strip().splitlines()
    added = set()
    for kw in still_missing:
        inserted = False
        for i, line in enumerate(lines):
            if ":" in line:
                cat_lower = line.split(":")[0].lower()
                kw_lower = kw.lower()
                if any(
                    tag in cat_lower
                    for tag in _keyword_category_hints(kw_lower)
                ):
                    lines[i] = line.rstrip().rstrip(",") + ", " + kw
                    added.add(kw)
                    inserted = True
                    break
        if not inserted and kw not in added:
            if lines:
                last = lines[-1]
                if ":" in last:
                    lines[-1] = last.rstrip().rstrip(",") + ", " + kw
                else:
                    lines.append(f"Additional Skills: {kw}")
            added.add(kw)
    return "\n".join(lines)


def _keyword_category_hints(kw: str) -> List[str]:
    """Return category name hints for a keyword to help categorize it."""
    kw = kw.lower()
    if any(t in kw for t in ["python", "java", "go ", "c++", "c/c++", "sql", "javascript"]):
        return ["language"]
    if any(t in kw for t in ["llm", "gpt", "genai", "rag", "prompt", "agent", "lam", "openai",
                               "gemini", "llama", "claude", "qwen", "foundational model"]):
        return ["generative", "llm", "ai"]
    if any(t in kw for t in ["ml", "machine learning", "deep learning", "model", "neural",
                               "classification", "regression", "anomaly", "predictive",
                               "time-series", "forecasting", "statistical"]):
        return ["ml", "dl", "machine"]
    if any(t in kw for t in ["aws", "s3", "ec2", "lambda", "sagemaker", "ecs", "eks",
                               "terraform", "cloud", "docker", "container", "serverless"]):
        return ["cloud", "infrastructure"]
    if any(t in kw for t in ["pandas", "numpy", "fastapi", "streamlit", "flask", "django",
                               "pytorch", "tensorflow", "scikit", "git", "docker"]):
        return ["framework", "tool"]
    if any(t in kw for t in ["ownership", "communication", "collaboration", "leadership",
                               "problem-solving", "teamwork", "urgency", "mentoring",
                               "analytical", "cross-functional"]):
        return ["soft"]
    return ["additional", "framework", "tool"]


def assemble_resume(
    original: str,
    summary: Dict,
    skills: Dict,
    projects: Dict,
    relevant_exp: Dict | None = None,
    job_title: str = "",
    missing_keywords: List[str] | None = None,
) -> str:
    name, contact = extract_name_contact(original)
    contact = re.sub(r"\|\s*(?=\S)", "| ", contact or "")
    summ = normalize(summary.get("optimized_summary", ""))

    sk = skills.get("skills_lines")
    if isinstance(sk, str):
        skills_block = normalize(sk)
    elif isinstance(sk, list):
        skills_block = normalize("\n".join(str(x) for x in sk))
    else:
        skills_block = ""

    # ── Relevant Experience ──────────────────────────────────────────
    rel_text = ""
    if relevant_exp:
        rel_text = normalize(relevant_exp.get("relevant_exp_text", ""))
        rel_text = fix_project_bullets(rel_text)
    if not rel_text:
        orig_rel = extract_relevant_experience(original)
        if orig_rel:
            for hdr in RELEVANT_EXP_HEADERS:
                orig_rel = re.sub(rf"(?mi)^\s*{re.escape(hdr)}\s*:?\s*$\n?", "", orig_rel)
            rel_text = normalize(orig_rel)

    # ── Projects ─────────────────────────────────────────────────────
    xp = normalize(projects.get("projects_text", ""))
    orig_proj = extract_projects(original)
    if "PROJECT" not in xp.upper() and orig_proj:
        xp = normalize(orig_proj)
    xp = re.sub(r"(?mi)^\s*PROJECTS\s*:?\s*$\n?", "", xp)
    xp = fix_project_bullets(xp)
    xp = dedupe_projects(xp)

    edu = extract_education(original) or ""

    parts = [
        name,
        contact or "[City, ST] | [Phone] | [Email] | LinkedIn | GitHub",
    ]

    parts.extend([
        "",
        "PROFESSIONAL SUMMARY",
        summ,
        "",
        "TECHNICAL SKILLS",
        skills_block,
    ])

    if rel_text:
        parts.extend(["", "RELEVANT EXPERIENCE", rel_text])

    parts.extend(["", "PROJECTS", xp])

    if edu:
        parts.extend(["", edu.strip()])
    return normalize("\n".join(parts))


# ── Pipeline (sequential to stay within Groq TPM limits) ─────────────────────

DELAY_BETWEEN_CALLS = 3  # seconds between LLM calls to avoid TPM throttling


def _make_groq_llm(model: str) -> ChatGroq:
    return ChatGroq(
        model=model,
        api_key=GROQ_API_KEY,
        temperature=0.0,
        max_tokens=4096,
    )


def _invoke_chain(prompt_template: str, llm, parser, variables: dict) -> str:
    chain = ChatPromptTemplate.from_template(prompt_template) | llm | parser
    result = chain.invoke(variables)
    time.sleep(DELAY_BETWEEN_CALLS)
    return result


def run_pipeline(
    resume: str,
    jd: str,
    *,
    extraction_model: str | None = None,
    assembler_model: str | None = None,
) -> Dict:
    if not GROQ_API_KEY:
        raise ValueError("GROQ_API_KEY is not set.")
    ex_id = (extraction_model or GROQ_EXTRACTION_MODEL).strip()
    asm_id = (assembler_model or GROQ_ASSEMBLER_MODEL).strip()
    llm_extract = _make_groq_llm(ex_id)
    llm_assemble = _make_groq_llm(asm_id)
    parser = StrOutputParser()
    base = {"resume": resume, "jd": jd}

    # ── Step 1: Keyword analysis ─────────────────────────────────────
    kw_raw = _invoke_chain(PROMPT_KEYWORDS, llm_extract, parser, base)
    kw = safe_json(kw_raw)

    # ── Step 2: Skills gap (uses same inputs) ────────────────────────
    _invoke_chain(PROMPT_SKILLS_GAP, llm_extract, parser, base)

    job_title = kw.get("job_title", "")
    for_summary = ", ".join(kw.get("for_summary", kw.get("jd_keywords", [])[:15]))
    for_skills = ", ".join(kw.get("for_skills", kw.get("jd_keywords", [])))
    for_bullets = ", ".join(kw.get("for_bullets", kw.get("jd_keywords", [])))
    missing_kws = ", ".join(kw.get("missing", []))
    missing_list = kw.get("missing", [])
    w2_vars = {
        **base,
        "job_title": job_title,
        "for_summary": for_summary,
        "for_skills": for_skills,
        "for_bullets": for_bullets,
        "missing_keywords": missing_kws,
    }

    # ── Step 3: Summary ──────────────────────────────────────────────
    sm = safe_json(_invoke_chain(PROMPT_SUMMARY, llm_assemble, parser, w2_vars))

    # ── Step 4: Skills ───────────────────────────────────────────────
    sk = safe_json(_invoke_chain(PROMPT_SKILLS, llm_assemble, parser, w2_vars))

    # ── Step 5: Relevant Experience ──────────────────────────────────
    re_exp = safe_json(_invoke_chain(PROMPT_RELEVANT_EXP, llm_assemble, parser, w2_vars))

    # ── Step 6: Projects ─────────────────────────────────────────────
    pr = safe_json(_invoke_chain(PROMPT_PROJECTS, llm_assemble, parser, w2_vars))

    resume_text = assemble_resume(
        resume, sm, sk, pr,
        relevant_exp=re_exp,
        job_title=job_title,
        missing_keywords=missing_list,
    )

    # ── Step 7: Cold Email ───────────────────────────────────────────
    w3_vars = {"optimized_resume": resume_text, "jd": jd}
    em = safe_json(_invoke_chain(PROMPT_EMAIL, llm_assemble, parser, w3_vars))

    # ── Step 8: Cover Letter ─────────────────────────────────────────
    cv = safe_json(_invoke_chain(PROMPT_COVER, llm_assemble, parser, w3_vars))

    return {
        "resume_text": resume_text,
        "email_data": em,
        "cover_data": cv,
    }


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _add_border(para, color="000000", size=6):
    pPr = para._p.get_or_add_pPr()
    old = pPr.find(qn("w:pBdr"))
    if old is not None:
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(size))
    b.set(qn("w:space"), "1")
    b.set(qn("w:color"), color)
    pBdr.append(b)
    pPr.append(pBdr)


def _hyperlink(paragraph, text: str, url: str, font_size=11.0):
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    for tag, val in [("w:u", "single"), ("w:color", "0563C1")]:
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        rPr.append(el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)
    rFont = OxmlElement("w:rFonts")
    rFont.set(qn("w:ascii"), "Calibri")
    rFont.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFont)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)


def _run(paragraph, text, size=11, bold=False, color=C_BLACK, italic=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.font.color.rgb = color
    return r


# Cover letter / cold-email: http(s), www., emails, LinkedIn & GitHub words
_URL_IN_TEXT_RE = re.compile(r"https?://[^\s\)\]\>\"\']+|www\.[^\s\)\]\>\"\']+", re.I)
_EMAIL_IN_TEXT_RE = re.compile(r"[\w\.\-]+@[\w\.\-]+\.\w+")
_LI_GH_WORD_RE = re.compile(r"\b(LinkedIn|linkedin|GitHub|github)\b")


def _trim_trailing_url_punct(url: str) -> str:
    u = url.rstrip()
    while u and u[-1] in ".,;:!?\"')\\]":
        u = u[:-1]
    return u


def _link_fragments(text: str, lm: Dict[str, str]) -> List[Tuple[str, str]]:
    """Split text into (display, href). Empty href means plain text."""
    if not text:
        return []
    out: List[Tuple[str, str]] = []
    i = 0
    n = len(text)
    while i < n:
        m_url = _URL_IN_TEXT_RE.search(text, i)
        m_em = _EMAIL_IN_TEXT_RE.search(text, i)
        m_wd = _LI_GH_WORD_RE.search(text, i)
        candidates: List[Tuple[int, str, re.Match]] = []
        if m_url:
            candidates.append((m_url.start(), "url", m_url))
        if m_em:
            candidates.append((m_em.start(), "em", m_em))
        if m_wd:
            candidates.append((m_wd.start(), "wd", m_wd))
        if not candidates:
            chunk = text[i:]
            if chunk:
                out.append((chunk, ""))
            break
        start, kind, m = min(candidates, key=lambda x: x[0])
        if start > i:
            out.append((text[i:start], ""))
        if kind == "url":
            raw = m.group(0)
            trimmed = _trim_trailing_url_punct(raw)
            href = trimmed
            if href.lower().startswith("www."):
                href = "https://" + href
            if trimmed:
                out.append((trimmed, href))
            i = m.end()
        elif kind == "em":
            addr = m.group(0)
            mail = lm.get("email", addr)
            out.append((addr, f"mailto:{mail}"))
            i = m.end()
        else:
            if m.group(1).lower() == "linkedin":
                out.append(("LinkedIn", lm["linkedin"]))
            else:
                out.append(("GitHub", lm["github"]))
            i = m.end()
    return out


def _add_paragraph_with_hyperlinks(
    paragraph,
    text: str,
    font_size: float,
    lm: Dict[str, str],
) -> None:
    for disp, href in _link_fragments(text, lm):
        if not disp:
            continue
        if href:
            _hyperlink(paragraph, disp, href, font_size)
        else:
            _run(paragraph, disp, size=font_size, color=C_BLACK)


def _is_section(line: str) -> bool:
    return line.upper().strip().rstrip(":") in SECTION_HEADERS


def _is_project_header(line: str) -> bool:
    return "|" in line and not _is_section(line)


def _is_bullet(line: str) -> bool:
    return bool(re.match(r"^[•·\-*▪◦]", line))


def _split_project_header(line: str) -> Tuple[str, str]:
    m = DATE_RE.search(line)
    if m:
        title = line[:m.start()].strip().rstrip("|").strip()
        date = m.group(1).strip()
        return title, date
    return line, ""


def _resolve_project_url(proj_name: str, lm: Dict[str, str]) -> str:
    key = proj_name.lower().strip()
    if key in lm:
        return lm[key]
    for lm_key, url in lm.items():
        if lm_key in key or key in lm_key:
            return url
    tokens = set(re.findall(r"[a-z]{3,}", key))
    if len(tokens) >= 2:
        for lm_key, url in lm.items():
            lm_tokens = set(re.findall(r"[a-z]{3,}", lm_key))
            if tokens & lm_tokens and len(tokens & lm_tokens) >= 2:
                return url
    return ""


def _unprotect(doc):
    settings = doc.settings.element
    for tag in ("w:documentProtection", "w:writeProtection"):
        for el in settings.findall(qn(tag)):
            settings.remove(el)


# ── DOCX builders ─────────────────────────────────────────────────────────────

# Resume .docx typography (pt) — name from first line of resume_text; contact row 12 pt;
# section headings 11.5 pt; body (bullets, projects, skills, education, etc.) 11 pt.
RESUME_NAME_PT = 18
RESUME_CONTACT_PT = 12
RESUME_SECTION_HEADING_PT = 11.5
RESUME_BODY_PT = 11


def build_resume_docx(resume_text: str) -> bytes:
    lm = LINK_MAP
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.5)
        sec.left_margin = sec.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(RESUME_BODY_PT)

    lines = [l for l in resume_text.split("\n") if l.strip()]
    name_done = False
    contact_done = False
    current_section = ""

    for line in lines:
        line = line.strip()

        # ── Name ──────────────────────────────────────────────────────────
        if not name_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, line, size=RESUME_NAME_PT, bold=True, color=C_BLACK)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            name_done = True
            continue

        # ── Contact line ──────────────────────────────────────────────────
        if not contact_done and ("|" in line or "@" in line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parts = [x.strip() for x in line.split("|")]
            for idx, part in enumerate(parts):
                email_m = re.search(r"[\w\.\-]+@[\w\.\-]+\.\w+", part)
                low = part.lower().strip()
                if email_m:
                    _hyperlink(
                        p,
                        part,
                        f"mailto:{lm.get('email', email_m.group(0))}",
                        RESUME_CONTACT_PT,
                    )
                elif "linkedin" in low:
                    _hyperlink(p, "LinkedIn", lm["linkedin"], RESUME_CONTACT_PT)
                elif "github" in low:
                    _hyperlink(p, "GitHub", lm["github"], RESUME_CONTACT_PT)
                else:
                    _run(p, part, size=RESUME_CONTACT_PT, color=C_GRAY)
                if idx < len(parts) - 1:
                    _run(p, " | ", size=RESUME_CONTACT_PT, color=C_GRAY)
            p.paragraph_format.space_after = Pt(2)
            hr = doc.add_paragraph()
            _add_border(hr, "000000", 8)
            hr.paragraph_format.space_before = Pt(0)
            hr.paragraph_format.space_after = Pt(2)
            contact_done = True
            continue

        # ── Section heading ───────────────────────────────────────────────
        if _is_section(line):
            current_section = line.upper().strip().rstrip(":")
            p = doc.add_paragraph()
            _run(p, current_section, size=RESUME_SECTION_HEADING_PT, bold=True, color=C_BLACK)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            _add_border(p, "000000", 6)
            continue

        in_projects = current_section in PROJECT_HEADERS
        in_experience = current_section in EXPERIENCE_HEADERS
        in_relevant_exp = current_section in RELEVANT_EXP_HEADERS
        in_skills = current_section in {
            "TECHNICAL SKILLS", "SKILLS", "CORE SKILLS", "KEY SKILLS",
        }
        in_education = current_section in {"EDUCATION", "ACADEMIC BACKGROUND"}

        # ── Bullet point ─────────────────────────────────────────────────
        if _is_bullet(line):
            p = doc.add_paragraph()
            txt = re.sub(r"^[•·\-*▪◦]\s*", "", line)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(0.25))
            _run(p, "•\t" + txt, size=RESUME_BODY_PT, color=C_BLACK)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            continue

        # ── Project header (Name | Tech   Date) ─────────────────────────
        if in_projects and _is_project_header(line):
            title_part, date_part = _split_project_header(line)
            p = doc.add_paragraph()

            pipe_parts = [x.strip() for x in title_part.split("|")]
            proj_name_raw = pipe_parts[0].strip() if pipe_parts else ""
            proj_url = _resolve_project_url(proj_name_raw, lm)

            for pidx, pp in enumerate(pipe_parts):
                low = pp.lower().strip()
                is_first = pidx == 0
                if (low == "link" or low.startswith("link")) and proj_url:
                    _hyperlink(p, pp, proj_url, RESUME_BODY_PT)
                else:
                    _run(p, pp, size=RESUME_BODY_PT, bold=is_first, color=C_BLACK)
                if pidx < len(pipe_parts) - 1:
                    _run(p, " | ", size=RESUME_BODY_PT, bold=False, color=C_BLACK)

            if date_part:
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(
                    Inches(6.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
                )
                _run(p, "\t", size=RESUME_BODY_PT)
                _run(p, date_part, size=RESUME_BODY_PT, bold=False, color=C_BLACK)

            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            continue

        # ── Experience / Relevant Experience entry header (Title | Org   Date)
        if (in_experience or in_relevant_exp) and _is_project_header(line):
            title_part, date_part = _split_project_header(line)
            p = doc.add_paragraph()

            pipe_parts = [x.strip() for x in title_part.split("|")]
            for pidx, pp in enumerate(pipe_parts):
                is_first = pidx == 0
                _run(p, pp, size=RESUME_BODY_PT, bold=is_first, color=C_BLACK)
                if pidx < len(pipe_parts) - 1:
                    _run(p, " | ", size=RESUME_BODY_PT, bold=False, color=C_BLACK)

            if date_part:
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(
                    Inches(6.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
                )
                _run(p, "\t", size=RESUME_BODY_PT)
                _run(p, date_part, size=RESUME_BODY_PT, bold=False, color=C_BLACK)

            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            continue

        # ── Skill category row (only in TECHNICAL SKILLS) ─────────────────
        if in_skills and ":" in line and len(line) < 500:
            colon = line.index(":")
            cat = line[:colon].strip()
            items = line[colon + 1:].strip()
            if cat and items:
                p = doc.add_paragraph()
                _run(p, cat + ":  ", size=RESUME_BODY_PT, bold=True, color=C_BLACK)
                _run(p, items, size=RESUME_BODY_PT, color=C_DARK)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                continue

        # ── Education lines (structured formatting) ───────────────────────
        if in_education:
            has_date = DATE_RE.search(line)
            is_inst = _is_institution_line(line)

            p = doc.add_paragraph()

            if has_date:
                text_part = line[:has_date.start()].strip().rstrip("|").strip()
                date_part = has_date.group(1).strip()
                _run(p, text_part, size=RESUME_BODY_PT, bold=is_inst, color=C_BLACK)
                tab_stops = p.paragraph_format.tab_stops
                tab_stops.add_tab_stop(
                    Inches(6.6), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
                )
                _run(p, "\t", size=RESUME_BODY_PT)
                _run(p, date_part, size=RESUME_BODY_PT, bold=False, color=C_BLACK)
            else:
                _run(p, line, size=RESUME_BODY_PT, bold=is_inst, color=C_BLACK)

            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(1)
            continue

        # ── Fallback — never bold ─────────────────────────────────────────
        p = doc.add_paragraph()
        _run(p, line, size=RESUME_BODY_PT, bold=False, color=C_BLACK)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(1)

    _unprotect(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_cover_docx(cover_data: dict) -> bytes:
    lm = LINK_MAP
    text = normalize(cover_data.get("cover_letter_text", ""))
    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(1.0)
        sec.left_margin = sec.right_margin = Inches(1.0)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    for para_text in [p.strip() for p in text.split("\n") if p.strip()]:
        p = doc.add_paragraph()
        _add_paragraph_with_hyperlinks(p, para_text, 11.0, lm)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(14)
    _unprotect(doc)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_email_txt(email_data: dict) -> bytes:
    subj = email_data.get("subject_line", "").strip()
    body = normalize(email_data.get("email_body", ""))
    return f"SUBJECT: {subj}\n\n{body}".encode("utf-8")


def build_zip(r: bytes, c: bytes, e: bytes, ts: str) -> bytes:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"ATS_Resume_{ts}.docx", r)
        zf.writestr(f"Cover_Letter_{ts}.docx", c)
        zf.writestr(f"Cold_Email_{ts}.txt", e)
    return buf.getvalue()


# ── Streamlit UI ──────────────────────────────────────────────────────────────

st.set_page_config(page_title="ATS Resume Builder v2", page_icon="📄", layout="wide")
st.markdown("## 📄 ATS Resume Builder v2")

with st.sidebar:
    st.subheader("Groq models")
    st.caption(
        "**Extraction** — ATS keywords + skills gap. **Assembler** — summary, skills, "
        "experience, projects, cold email, cover letter."
    )
    _ex_opts = _groq_select_options(GROQ_EXTRACTION_MODEL)
    _asm_opts = _groq_select_options(GROQ_ASSEMBLER_MODEL)
    extraction_model_ui = st.selectbox(
        "Extraction model",
        _ex_opts,
        index=_groq_select_index(_ex_opts, GROQ_EXTRACTION_MODEL),
    )
    assembler_model_ui = st.selectbox(
        "Assembler model",
        _asm_opts,
        index=_groq_select_index(_asm_opts, GROQ_ASSEMBLER_MODEL),
    )

if not GROQ_API_KEY:
    st.error("Set **GROQ_API_KEY** in your `.env` or environment variables.")

c1, c2 = st.columns(2)
with c1:
    up = st.file_uploader("Upload current resume (.docx)", type=["docx"])
with c2:
    jd = st.text_area("Job Description", height=260, placeholder="Paste the full JD here…")

go = st.button(
    "🚀 Generate ATS Resume + Cover Letter + Cold Email",
    type="primary",
    disabled=not (up and jd.strip() and GROQ_API_KEY),
)

if go and up and jd.strip():
    with st.spinner("Running 8 steps (keywords → summary → skills → relevant exp → projects → email → cover)… ~2 min"):
        up.seek(0)
        resume_plain = extract_text_from_docx(up)
        try:
            out = run_pipeline(
                resume_plain,
                jd.strip(),
                extraction_model=extraction_model_ui,
                assembler_model=assembler_model_ui,
            )
        except Exception as e:
            st.error(str(e))
            st.stop()

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    r_bytes = build_resume_docx(out["resume_text"])
    cv_bytes = build_cover_docx(out["cover_data"])
    em_bytes = build_email_txt(out["email_data"])
    zip_bytes = build_zip(r_bytes, cv_bytes, em_bytes, ts)

    st.success("Done! Files are fully editable — if Word shows 'Protected View', click **Enable Editing**.")

    st.download_button(
        "⬇️  Download All (.zip)",
        data=zip_bytes,
        file_name=f"ATS_Package_{ts}.zip",
        mime="application/zip",
        use_container_width=True,
        type="primary",
    )

    d1, d2, d3 = st.columns(3)
    with d1:
        st.download_button(
            "⬇️ Resume (.docx)",
            data=r_bytes,
            file_name=f"ATS_Resume_{ts}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with d2:
        st.download_button(
            "⬇️ Cover Letter (.docx)",
            data=cv_bytes,
            file_name=f"Cover_Letter_{ts}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    with d3:
        st.download_button(
            "⬇️ Cold Email (.txt)",
            data=em_bytes,
            file_name=f"Cold_Email_{ts}.txt",
            mime="text/plain",
            use_container_width=True,
        )
